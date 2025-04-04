# to deploy in terminal
# rsconnect add --account habadio --name habadio --token 8D654CD9CCCEBE179F6DB7BE6764D040 --secret No1MZ0K+IHtQPDcwd7fNbJliPUVvvlcL5sJuuR3v
# rsconnect deploy shiny /Users/helenabadiotakis/Downloads/ChanLab/shiny_manuscript --name habadio --title shiny-manuscript


# https://metaboanalyst.ca/
# check iqr decimal placess
# option to add total as another group 
# ttest shapiro wilk test - normality test

"""
default no columns checked, 
step 1: upload file, show exa
step 2: select columns to include. only when check does it create a card to name/customize
step 3: input table_name, define subheadings
step 4: move up customization, dec places (one for p-values- default to 3 & table values), output format
step 5: customize variables in card, add option for grouping here-> 
    see if can include under subheadings and select + button to add
embed pic of variable options 

TernTables - Medical Research Table Creator
"""
# step 2: checkboxes for which columns to include in the table

# imports
from shiny import App, reactive, render, ui
import pandas as pd
import shinywidgets as sw
import os
from io import StringIO
import re
from scipy import stats
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pickle
from itertools import combinations
# import rpy2.robjects.numpy2ri
# from rpy2.robjects.packages import importr

# set default and alternative statistical tests
default_tests = {
    "Omit": "Omit",
    "Categorical (Y/N)": "fisher",
    "Categorical (Dichotomous)": "fisher",
    "Categorical (Multinomial)": "fisher-freeman-halton",
    "Ratio Continuous": "ttest",
    "Ordinal Discrete": "wilcoxon",
}

alternative_tests = {
    "Categorical (Y/N)": "chi2",
    "Categorical (Dichotomous)": "chi2",
    "Categorical (Multinomial)": "chi2",
    "Ratio Continuous": "mannwhitney",
    "Ordinal Discrete": "ttest",
}

variable_types = list(default_tests.keys())

file_path = '/Users/helenabadiotakis/Downloads/misc_materials/example_data_test.csv'
df = pd.read_csv(file_path)  # Reads header row by default

columns = df.columns.tolist()  # Get column names
columns = [re.sub(r'\W+', '', col) for col in columns]

subheadings = ["Demographics", "Donor", "Other"]

# Store variable settings in a dictionary
var_config = {}
var_config["group"] = {
    "type": "Omit", 
    "rename": "group", 
    "subheading": "Other",
    "position": 4,
    "p_value": None
}
var_config["age"] = {
    "type": "Ratio Continuous", 
    "rename": "Age", 
    "subheading": "Demographics",
    "position": 1,
    "p_value": None
}
var_config["race"] = {
    "type": "Categorical (Multinomial)", 
    "rename": "Race", 
    "subheading": "",
    "position": 0,
    "p_value": None
}
var_config["sex"] = {
    "type": "Categorical (Dichotomous)", 
    "rename": "Sex", 
    "subheading": "Demographics",
    "position": 2,
    "p_value": None
}
var_config["bmi"] = {
    "type": "Ordinal Discrete", 
    "rename": "BMI", 
    "subheading": "Demographics",
    "position": 3,
    "p_value": None
}
var_config["cig_use"] = {
    "type": "Categorical (Y/N)", 
    "rename": "History of Smoking", 
    "subheading": "Donor",
    "position": 1,
    "p_value": None
}
var_config["diab"] = {
    "type": "Categorical (Y/N)", 
    "rename": "DM", 
    "subheading": "Demographics",
    "position": 2,
    "p_value": None
}
var_config["mcs"] = {
    "type": "Categorical (Multinomial)", 
    "rename": "mcs", 
    "subheading": "Donor",
    "position": 3,
    "p_value": None
}
grouping_var = "group"

decimal_places = 2
# output_format = "n (%)"
output_format = "% (n)"

# get p-values from statistical test
################################################################################
### ONLY SUPPORTS 2 GROUPS AT THE MOMENT, NEED TO UPDATE TO MULTIPLE GROUPS ####
################################################################################
def run_statistical_test(df, group_var, var_type, var_name, decimal_places):
    groups = df[group_var].unique()
    if len(groups) != 2:
        return None  # Only supports two-group comparisons
    
    group1 = df[df[group_var] == groups[0]][var_name].dropna()
    group2 = df[df[group_var] == groups[1]][var_name].dropna()
    
    test_type = default_tests[var_type]

    if test_type == "fisher":
        contingency_table = pd.crosstab(df[var_name], df[group_var])
        _, p_value = stats.fisher_exact(contingency_table)
    elif test_type == 'fisher-freeman-halton':
        contingency_table = pd.crosstab(df[var_name], df[group_var])
        _, p_value, _, _ = stats.chi2_contingency(contingency_table, lambda_="log-likelihood")        
        # rpy2.robjects.numpy2ri.activate()
        # stats = importr('stats')
        # m = np.array([[4,4],[4,5],[10,6]])
        # res = stats.fisher_test(m)
        # print 'p-value: {}'.format(res[0][0])
    elif test_type == "chi2":
        contingency_table = pd.crosstab(df[var_name], df[group_var])
        _, p_value, _, _ = stats.chi2_contingency(contingency_table)
    elif test_type == "ttest":
        _, p_value = stats.ttest_ind(group1, group2, equal_var=False)
    elif test_type == "mannwhitney":
        _, p_value = stats.mannwhitneyu(group1, group2)
    elif test_type == "wilcoxon":
        _, p_value = stats.ranksums(group1, group2)
    else:
        print("Invalid test type:", test_type, " for variable:", var_name)
        p_value = None
    
    try:
        p_value = round(p_value, decimal_places)
    except:
        pass

    return p_value

# Function to perform aggregation analysis based on the variable type
def perform_aggregate_analysis(df, group_var, var_type, var_name, decimal_places, output_format, col_var_config):
    groups = df[group_var].unique()
    if len(groups) != 2:
        return None  # Only supports two-group comparisons
    # print(group_var, test_type, var_name, decimal_places, output_format, col_var_config)
    
    # Check if the variable has a "Yes" option
    yes_values = ['Yes', 'Y', 'y', 'yes', 1]
    yn_var = None

    var_options = df[var_name].unique()        
    for val in yes_values:
        if val in var_options:
            yn_var=val

    group1 = df[df[group_var] == groups[0]][var_name].dropna()
    group2 = df[df[group_var] == groups[1]][var_name].dropna()

    group1_total = len(group1)
    group2_total = len(group2)

    # Default result structure
    result = {}
    
    if var_type == "Omit":
        return None
    
    elif var_type == "Categorical (Y/N)":
        # count occurrences of yn_var in group 1 and group 2
        group1_sum = (group1 == yn_var).sum()
        group2_sum = (group2 == yn_var).sum()

        # Store the aggregate values in var_config
        if output_format == "n (%)":
            col_var_config['group1'] = str(group2_sum)  + " (" + str(round(group2_sum / group2_total * 100, decimal_places)) + "%)"
            col_var_config['group2'] = str(group1_sum) + " (" + str(round(group1_sum / group1_total * 100, decimal_places)) + "%)"
        else:
            col_var_config['group1'] = str(round(group1_sum / group1_total * 100, decimal_places)) + "% (" + str(group1_sum) + ")"
            col_var_config['group2'] = str(round(group2_sum / group2_total * 100, decimal_places)) + "% (" + str(group2_sum) + ")"
        
    elif var_type == "Categorical (Dichotomous)" or var_type == "Categorical (Multinomial)":
        for i in range(len(var_options)):
            group1_sum = (group1 == var_options[i]).sum()
            group2_sum = (group2 == var_options[i]).sum()

            if output_format == "n (%)":
                col_var_config[f'group1_subgroup{i}'] = str(group1_sum) + " (" + str(round(group1_sum / group1_total * 100, decimal_places)) + "%)"
                col_var_config[f'group2_subgroup{i}'] = str(group2_sum)  + " (" + str(round(group2_sum / group2_total * 100, decimal_places)) + "%)"
            else:
                col_var_config[f'group1_subgroup{i}'] = str(round(group1_sum / group1_total * 100, decimal_places)) + "% (" + str(group1_sum) + ")"
                col_var_config[f'group2_subgroup{i}'] = str(round(group2_sum / group2_total * 100, decimal_places)) + "% (" + str(group2_sum) + ")"

    elif var_type == "Ratio Continuous":
        # Aggregate: Mean and Standard Deviation
        group1_mean = round(group1.mean(), decimal_places)
        group2_mean = round(group2.mean(), decimal_places)

        group1_std = round(group1.std(), decimal_places)
        group2_std = round(group2.std(), decimal_places)

        col_var_config['group1'] = str(group1_mean) + " \u00B1 " + str(group1_std)
        col_var_config['group2'] = str(group2_mean) + " \u00B1 " + str(group2_std)

    elif var_type == "Ordinal Discrete":
        # Aggregate: Median and Interquartile Range (IQR)
        group1_median = group1.median()
        group2_median = group2.median()
        group1_iqr = [group1.quantile(0.25), group1.quantile(0.75)]
        group2_iqr = [group2.quantile(0.25), group2.quantile(0.75)]

        col_var_config['group1'] = str(group1_median) + " [" + str(group1_iqr[0]) + "-" + str(group1_iqr[1]) + "]"
        col_var_config['group2'] = str(group2_median) + " [" + str(group2_iqr[0]) + "-" + str(group2_iqr[1]) + "]"
    
    return col_var_config

# Function to create Word table from var_config
def create_word_table(df,var_config, subheadings):
    # Create a new Word Document
    doc = Document()

    # Create the table with columns for Variable, Group 1, Group 2, P-Value
    table = doc.add_table(rows=1, cols=4)
    table.columns[0].width=Inches(3.5)
    table.columns[1].width=Inches(1.5)
    table.columns[2].width=Inches(1.5)
    table.columns[3].width=Inches(.5)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Variable'
    hdr_cells[1].text = 'Group 1'
    hdr_cells[2].text = 'Group 2'
    hdr_cells[3].text = 'P-Value'
    
    for row in hdr_cells:
        row.paragraphs[0].runs[0].font.bold = True  # Bold formatting for the subheading row
        
    group1_size = len(df[df[grouping_var] == df[grouping_var].unique()[0]])
    group2_size = len(df[df[grouping_var] == df[grouping_var].unique()[1]])
    grp_cells = table.add_row().cells
    grp_cells[0].text = ''
    grp_cells[1].text = '(n= ' + str(group1_size) + ")"
    grp_cells[2].text = '(n= ' + str(group2_size) + ")"
    grp_cells[3].text = ''

    # Loop through subheadings
    for subheading_name in subheadings:
        # Add a row for the subheading (this is the row header)
        row_cells = table.add_row().cells
        row_cells[0].text = f"{subheading_name}"  # Subheading name in the first column
        row_cells[1].text = ''  # Leave empty for Group 1
        row_cells[2].text = ''  # Leave empty for Group 2
        row_cells[3].text = ''  # Leave empty for P-Value

        row_cells[0].paragraphs[0].runs[0].font.bold = True  # Bold formatting for the subheading row
        
        # Get and sort all variables for the current subheading
        subheading_vars = [col for col, config in var_config.items() if config['subheading'] == subheading_name]
        sorted_subheading_vars = sorted(subheading_vars, key=lambda x: var_config[x]["position"])
        
        # Add a row for each variable under the current subheading
        for var in sorted_subheading_vars:
            var_type = var_config[var]["type"]
            var_name = var_config[var]["rename"]

            if var_type == "Omit":
                continue
            elif var_type == "Categorical (Y/N)":
                row_cells = table.add_row().cells
                row_cells[0].text = f"{var_name}"  
                row_cells[1].text = str(var_config[var]["group1"])
                row_cells[2].text = str(var_config[var]["group2"])
                row_cells[3].text = str(var_config[var]["p_value"])

            elif var_type == "Categorical (Dichotomous)":
                row_cells = table.add_row().cells
                row_cells[0].text = f"{var_name}"  
                row_cells[1].text = ""
                row_cells[2].text = ""
                row_cells[3].text = ""

                row_cells[0].paragraphs[0].runs[0].font.underline = True

                var_options = df[var].unique()        
                for i in range(len(var_options)):
                    row_cells = table.add_row().cells
                    row_cells[0].text = f"   {var_options[i]}"  
                    row_cells[1].text = str(var_config[var][f"group1_subgroup{i}"])
                    row_cells[2].text = str(var_config[var][f"group2_subgroup{i}"])
                    if i == 0:
                        row_cells[3].text = str(var_config[var]["p_value"])
                    else:
                        row_cells[3].text = "-"

                    row_cells[0].paragraphs[0].runs[0].font.italic = True

            elif var_type == "Categorical (Multinomial)":
                row_cells = table.add_row().cells
                row_cells[0].text = f"{var_name}"  
                row_cells[1].text = ""
                row_cells[2].text = ""
                row_cells[3].text = ""

                row_cells[0].paragraphs[0].runs[0].font.underline = True

                var_options = df[var].unique()        
                for i in range(len(var_options)):
                    row_cells = table.add_row().cells
                    row_cells[0].text = f"   {var_options[i]}"  
                    row_cells[1].text = str(var_config[var][f"group1_subgroup{i}"])
                    row_cells[2].text = str(var_config[var][f"group2_subgroup{i}"])
                    if i == 0:
                        row_cells[3].text = str(var_config[var]["p_value"])
                    else:
                        row_cells[3].text = "-"

                    row_cells[0].paragraphs[0].runs[0].font.italic = True

            elif var_type == "Ratio Continuous" or var_type == "Ordinal Discrete":
                row_cells = table.add_row().cells
                row_cells[0].text = f"{var_name}"  
                row_cells[1].text = str(var_config[var]["group1"])
                row_cells[2].text = str(var_config[var]["group2"])
                row_cells[3].text = str(var_config[var]["p_value"])
                
                # Apply formatting to the variable name cell (indentation and smaller font)
                # para = row_cells[0].paragraphs[0]
                # run = para.add_run(row_cells[0].text)
                # run.font.size = Pt(8)  # Smaller font size
                # para.paragraph_format.left_indent = Pt(12)  # Indentation for the variable name            

    # Save the document to a file
    doc_filename = "statistical_analysis_results.docx"
    doc.save(doc_filename)
    return doc_filename


# remove grouping_var from columns
columns.remove(grouping_var)
print(columns)
# Perform statistical analysis using the grouping variable
for col in columns:
    var_type = var_config[col]["type"]
    
    if var_type != "Omit":
        p_value = run_statistical_test(df, grouping_var, var_type, col, decimal_places)
        var_config[col]["p_value"] = p_value

        aggregate_result = perform_aggregate_analysis(df, grouping_var
                                                      , var_type, col, decimal_places, output_format, var_config[col])
        var_config[col] = aggregate_result
        # print(var_config[col]["p_value"], default_tests[var_type])
        print(var_config[col])

print(type(var_config),var_config.keys())
doc_filename = create_word_table(df, var_config, subheadings)
