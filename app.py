
# imports
from shiny import App, reactive, render, ui
import pandas as pd
import shinywidgets as sw
import os
from io import StringIO
import re
import numpy as np
from scipy import stats
from statsmodels.stats.contingency_tables import Table2x2
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
# from fisher import pvalue_nway
import pickle
import json

# set default and alternative statistical tests
default_tests = {
    "Omit": "Omit",
    "Categorical (Y/N)": "fisher",
    "Categorical (Dichotomous)": "fisher",
    "Categorical (Multinomial)": "fisher-freeman-halton",
    "Ratio Continuous": "ttest",
    "Ordinal Discrete": "wilcoxon",
}

alternative_tests = {
    "Categorical (Y/N)": "chi2",
    "Categorical (Dichotomous)": "chi2",
    "Categorical (Multinomial)": "chi2",
    "Ratio Continuous": "mannwhitney",
    "Ordinal Discrete": "ttest",
}

missing_values = ["NA", "N/A", "NAN", "na", "n/a", "nan", "Na", "unk", "unknown", "Unk", "Unknown", "UNKNOWN"] # List of strings representing unknown or missing data
 
# variable_types = [
#     "Omit",
#     "Binary (i.e. Smoking, Diabetes, Hypertension)",
#     "Categorical (Dichotomous) (i.e. Sex)",
#     "Categorical (Multinomial) (i.e. Race)",
#     "Ratio Continuous (i.e., Age, BMI)",
#     "Ordinal Discrete (i.e., GCS, Tumor Grade)",
# ]

variable_types = list(default_tests.keys())

# get p-values from statistical test
################################################################################
### ONLY SUPPORTS 2 GROUPS AT THE MOMENT, NEED TO UPDATE TO MULTIPLE GROUPS ####
################################################################################
def run_statistical_test(df, group_var, var_type, var_name, decimal_places):
    print("PERFORM PVAL TEST", group_var, var_type, var_name, df[var_name].dropna().unique())
    groups = df[group_var].dropna().unique()
    
    if len(groups) != 2:
        print("Only supports two-group comparisons")
        return None  # Only supports two-group comparisons

    group1 = df[df[group_var] == groups[0]][var_name].dropna()
    group2 = df[df[group_var] == groups[1]][var_name].dropna()

    test_type = default_tests[var_type]

    if test_type == "fisher":
        contingency_table = pd.crosstab(df[var_name], df[group_var])
        _, p_value = stats.fisher_exact(contingency_table)
    elif test_type == 'fisher-freeman-halton': # UPDATE
        contingency_table = pd.crosstab(df[var_name], df[group_var])
        _, p_value, _, _ = stats.chi2_contingency(contingency_table, lambda_="log-likelihood")        
        # p_value = pvalue_nway(contingency_table.values).two_tail
    elif test_type == "chi2":
        contingency_table = pd.crosstab(df[var_name], df[group_var])
        _, p_value, _, _ = stats.chi2_contingency(contingency_table)
    elif test_type == "ttest":
        _, p_value = stats.ttest_ind(group1, group2, equal_var=False)
    elif test_type == "mannwhitney":
        _, p_value = stats.mannwhitneyu(group1, group2)
    elif test_type == "wilcoxon":
        _, p_value = stats.ranksums(group1, group2)
    else:
        print("Invalid test type:", test_type, " for variable:", var_name)
        p_value = None
    
    try:
        p_value = round(p_value, decimal_places)
    except:
        pass

    return p_value

def compute_odds_ratio_between_groups(group1, group2, reference_value):
    """
    Calculates odds ratio and 95% CI between two groups for a binary categorical variable.

    Parameters:
    - group1: pd.Series of exposure variable values for group 1 (e.g., var_name from df[group_var] == Group1)
    - group2: pd.Series of exposure variable values for group 2
    - reference_value: The reference value (default "No")

    Returns:
    - String in format "OR [low–high]" or "undefined" if the table is invalid
    """
    # Drop NA
    group1 = group1.dropna().astype(str)
    group2 = group2.dropna().astype(str)

    # Convert to binary: 0 = reference, 1 = not reference
    group1_binary = (group1 != reference_value).astype(int)
    group2_binary = (group2 != reference_value).astype(int)

    # Build 2x2 contingency table:
    #         | not reference (1) | reference (0)
    # group1 |        a          |      b
    # group2 |        c          |      d
    a = (group1_binary == 1).sum()
    b = (group1_binary == 0).sum()
    c = (group2_binary == 1).sum()
    d = (group2_binary == 0).sum()

    table = [[a, b], [c, d]]
    
    print(group1_binary, group2_binary)
    print("Ref Value: ", reference_value )
    print(table)

    # Ensure valid 2x2 table
    if any(v == 0 for row in table for v in row):
        return "undefined"

    try:
        ct = Table2x2(table)
        or_value = ct.oddsratio
        ci_low, ci_high = ct.oddsratio_confint()
        return f"{or_value:.2f} [{ci_low:.2f}–{ci_high:.2f}]"
    except Exception as e:
        return f"error: {str(e)}"

# Function to perform aggregation analysis based on the variable type
def perform_aggregate_analysis(df, group_var, var_type, var_name, decimal_places, output_format, col_var_config):
    print("PERFORM AGG ANALYSIS", group_var, var_type, var_name, df[var_name].dropna().unique(), decimal_places, output_format, col_var_config)
    groups = df[group_var].dropna().unique()
    if len(groups) != 2:
        print("Only supports two-group comparisons")
        return None  # Only supports two-group comparisons
    
    # Check if the variable has a "Yes" option
    yes_values = ['Yes', 'Y', 'y', 'yes', 1]
    yn_var = None

    var_options = df[var_name].dropna().unique()
    for val in yes_values:
        if val in var_options:
            yn_var=val
        
    group1 = df[df[group_var] == groups[0]][var_name].dropna()
    group2 = df[df[group_var] == groups[1]][var_name].dropna()

    group1_total = len(group1)
    group2_total = len(group2)
    
    if var_type == "Omit":
        return None
    
    elif var_type == "Categorical (Y/N)":
        # count occurrences of yn_var in group 1 and group 2
        group1_sum = (group1 == yn_var).sum()
        group2_sum = (group2 == yn_var).sum()

        # Store the aggregate values in var_config
        if output_format == "n (%)":
            col_var_config['group1'] = str(group2_sum)  + " (" + str(round(group2_sum / group2_total * 100, decimal_places)) + "%)"
            col_var_config['group2'] = str(group1_sum) + " (" + str(round(group1_sum / group1_total * 100, decimal_places)) + "%)"
        else:
            col_var_config['group1'] = str(round(group1_sum / group1_total * 100, decimal_places)) + "% (" + str(group1_sum) + ")"
            col_var_config['group2'] = str(round(group2_sum / group2_total * 100, decimal_places)) + "% (" + str(group2_sum) + ")"

        col_var_config['odds_ratio'] = compute_odds_ratio_between_groups(group1, group2, reference_value=col_var_config['ref_val'])
        
    elif var_type == "Categorical (Dichotomous)" or var_type == "Categorical (Multinomial)":
        print(var_options)
        for i in range(len(var_options)):
            group1_sum = (group1 == var_options[i]).sum()
            group2_sum = (group2 == var_options[i]).sum()

            if output_format == "n (%)":
                col_var_config[f'group1_subgroup{i}'] = str(group1_sum) + " (" + str(round(group1_sum / group1_total * 100, decimal_places)) + "%)"
                col_var_config[f'group2_subgroup{i}'] = str(group2_sum)  + " (" + str(round(group2_sum / group2_total * 100, decimal_places)) + "%)"
            else:
                col_var_config[f'group1_subgroup{i}'] = str(round(group1_sum / group1_total * 100, decimal_places)) + "% (" + str(group1_sum) + ")"
                col_var_config[f'group2_subgroup{i}'] = str(round(group2_sum / group2_total * 100, decimal_places)) + "% (" + str(group2_sum) + ")"
            
            col_var_config['odds_ratio'] = compute_odds_ratio_between_groups(group1, group2, reference_value=col_var_config['ref_val'])

    elif var_type == "Ratio Continuous":
        # Aggregate: Mean and Standard Deviation
        group1_mean = round(group1.mean(), decimal_places)
        group2_mean = round(group2.mean(), decimal_places)

        group1_std = round(group1.std(), decimal_places)
        group2_std = round(group2.std(), decimal_places)

        col_var_config['group1'] = str(group1_mean) + " \u00B1 " + str(group1_std)
        col_var_config['group2'] = str(group2_mean) + " \u00B1 " + str(group2_std)

    elif var_type == "Ordinal Discrete":
        # Aggregate: Median and Interquartile Range (IQR)
        group1_median = group1.median()
        group2_median = group2.median()
        group1_iqr = [group1.quantile(0.25), group1.quantile(0.75)]
        group2_iqr = [group2.quantile(0.25), group2.quantile(0.75)]

        col_var_config['group1'] = str(group1_median) + " [" + str(group1_iqr[0]) + "-" + str(group1_iqr[1]) + "]"
        col_var_config['group2'] = str(group2_median) + " [" + str(group2_iqr[0]) + "-" + str(group2_iqr[1]) + "]"
    
    return col_var_config

# Function to create Word table from var_config
def create_word_table(df,var_config, group_var, subheadings, subheading_names, table_name, odds_ratio, output_format):
    if odds_ratio == 'Yes':
        odds_ratio = True
    else:
        odds_ratio = False

    # Create a new Word Document
    doc = Document()

    # Create the table with columns for Variable, Group 1, Group 2, P-Value
    if odds_ratio:
        table = doc.add_table(rows=1, cols=5)
    else:
        table = doc.add_table(rows=1, cols=4)
    table.columns[0].width=Inches(3)
    table.columns[1].width=Inches(1.5)
    table.columns[2].width=Inches(1.5)
    table.columns[3].width=Inches(.5)
    if odds_ratio:
        table.columns[4].width=Inches(1.5)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Variable'
    hdr_cells[1].text = f'{df[group_var].unique()[0]}'
    hdr_cells[2].text = f'{df[group_var].unique()[1]}'
    hdr_cells[3].text = 'P-Value'
    if odds_ratio:
        hdr_cells[4].text='Odds Ratio'

    
    for row in hdr_cells:
        row.paragraphs[0].runs[0].font.bold = True  # Bold formatting for the subheading row
        
    group1_size = len(df[df[group_var] == df[group_var].unique()[0]])
    group2_size = len(df[df[group_var] == df[group_var].unique()[1]])
    grp_cells = table.add_row().cells
    grp_cells[0].text = ''
    grp_cells[1].text = '(n= ' + str(group1_size) + ")"
    grp_cells[2].text = '(n= ' + str(group2_size) + ")"
    grp_cells[3].text = ''
    if odds_ratio:
        grp_cells[4].text = ''


    # Loop through subheadings
    for sub in subheadings:
        subheading_name = subheading_names[sub]()

        # Add a row for the subheading (this is the row header)
        row_cells = table.add_row().cells
        row_cells[0].text = f"{subheading_name}"  # Subheading name in the first column
        row_cells[1].text = ''  # Leave empty for Group 1
        row_cells[2].text = ''  # Leave empty for Group 2
        row_cells[3].text = ''  # Leave empty for P-Value
        if odds_ratio:
            row_cells[4].text = ''  # Leave empty for Odds Ratio

        row_cells[0].paragraphs[0].runs[0].font.bold = True  # Bold formatting for the subheading row
        
        # Get and sort all variables for the current subheading
        subheading_vars = [col for col, config in var_config.items() if config["subheading"] == sub]
        subheading_vars = [col for col in subheading_vars if col != group_var]
        sorted_subheading_vars = sorted(subheading_vars, key=lambda x: var_config[x]["position"])
        
        # Add a row for each variable under the current subheading
        for var in sorted_subheading_vars:
            var_name = var_config[var]["name"]
            var_type = var_config[var]["type"]
            if var_type == "Omit":
                continue
            elif var_type == "Categorical (Y/N)":
                row_cells = table.add_row().cells
                row_cells[0].text = f"   {var_name}"  
                row_cells[1].text = str(var_config[var]["group1"])
                row_cells[2].text = str(var_config[var]["group2"])
                row_cells[3].text = str(var_config[var]["p_value"])
                if odds_ratio:
                    row_cells[4].text = str(var_config[var]["odds_ratio"])

            elif var_type == "Categorical (Dichotomous)":
                row_cells = table.add_row().cells
                row_cells[0].text = f"   {var_name}"  
                row_cells[1].text = ""
                row_cells[2].text = ""
                row_cells[3].text = ""
                
                # row_cells[0].paragraphs[0].runs[0].font.underline = True

                var_options = df[var].dropna().unique().tolist()
                ref_val = var_config[var]["ref_val"]
                if ref_val in var_options:
                    var_options.remove(ref_val)
                    var_options.insert(0, ref_val)
                
                for i in range(len(var_options)):
                    row_cells = table.add_row().cells
                    row_cells[0].text = f"      {var_options[i]}"  
                    row_cells[1].text = str(var_config[var][f"group1_subgroup{i}"])
                    row_cells[2].text = str(var_config[var][f"group2_subgroup{i}"])
                    if var_options[i] == var_config[var]["ref_val"]:
                        row_cells[3].text = str(var_config[var]["p_value"])
                    else:
                        row_cells[3].text = "-"
                    if var_options[i] == var_config[var]["ref_val"] and odds_ratio:
                        row_cells[4].text = str(var_config[var]["odds_ratio"])

                    row_cells[0].paragraphs[0].runs[0].font.italic = True

            elif var_type == "Categorical (Multinomial)":
                row_cells = table.add_row().cells
                row_cells[0].text = f"   {var_name}"  
                row_cells[1].text = ""
                row_cells[2].text = ""
                row_cells[3].text = ""

                # row_cells[0].paragraphs[0].runs[0].font.underline = True

                var_options = df[var].dropna().unique()        
                for i in range(len(var_options)):
                    row_cells = table.add_row().cells
                    row_cells[0].text = f"      {var_options[i]}"  
                    row_cells[1].text = str(var_config[var][f"group1_subgroup{i}"])
                    row_cells[2].text = str(var_config[var][f"group2_subgroup{i}"])
                    if i == 0:
                        row_cells[3].text = str(var_config[var]["p_value"])
                    else:
                        row_cells[3].text = "-"

                    row_cells[0].paragraphs[0].runs[0].font.italic = True

            elif var_type == "Ratio Continuous" or var_type == "Ordinal Discrete":
                row_cells = table.add_row().cells
                row_cells[0].text = f"   {var_name}"  
                row_cells[1].text = str(var_config[var]["group1"])
                row_cells[2].text = str(var_config[var]["group2"])
                row_cells[3].text = str(var_config[var]["p_value"])

    doc.add_page_break()  # Add a page break after the table

    # Add summary of statistical tests
    present_types = set(config["type"] for config in var_config.values())
    sentences = []

    if any(t in present_types for t in ["Categorical (Y/N)", "Categorical (Dichotomous)"]):
        test_name = "Fisher's exact test" #update when alternative stat test is added
        sentences.append(
            f"All dichotomous/binary variables were analyzed with {test_name} and are displayed as {output_format}."
        )

    if "Categorical (Multinomial)" in present_types:
        test_name = "Fisher-Freeman-Halton test" #update when alternative stat test is added
        sentences.append(
            f"All multinomial variables were analyzed with {test_name} and are displayed as {output_format}."
        )

    if "Ordinal Discrete" in present_types:
        sentences.append(
            "Ordinal discrete variables were analyzed using a Wilcoxon rank sum test and are displayed as median [interquartile range]."
        )

    if "Ratio Continuous" in present_types:
        sentences.append(
            "For continuous variables, normality tests were applied using the Shapiro-Wilk test. If a variable failed the normality test (Shapiro-Wilk p < 0.05), then a non-parametric test (Mann-Whitney U test) was used for analysis. If a variable passed the normality test (Shapiro-Wilk p > 0.05), then a student's t-test was performed. All continuous variables are analyzed via student's t-test and are displayed as mean ± standard deviation."
        )

    var_config_summary = " ".join(sentences)
    doc.add_paragraph(var_config_summary)  

    # Save the document to a file
    table_name = re.sub(r'\W+', '', table_name.strip())
    if table_name == "":
        table_name = "Statistical_Analysis"
    doc_filename = f"{table_name}.docx"
    doc.save(doc_filename)
    return doc_filename


# Save configuration to a .pkl file
def save_config(config, filename="config.pkl"):
    with open(filename, "wb") as f:
        pickle.dump(config, f)
    print(f"Configuration saved as {filename}")

# Load configuration from a .pkl file
def load_config(filename="config.pkl"):
    try:
        with open(filename, "rb") as f:
            config = pickle.load(f)
        return config
    except FileNotFoundError:
        print("No saved configuration found.")
        return None
    


################################################################################
######################### Shiny App Layout #####################################
################################################################################
app_ui = ui.page_fluid(
    # Custom CSS for styling
    ui.tags.style("""
            .section-title {
                font-weight: bold;
                text-align: center;
                font-size: 36px;
                margin-top: 20px;
                margin-bottom: 30px;
            }
            .step-header {
                font-weight: bold;
                margin-top: 50px;
                margin-bottom: 10px;
            }
            .variable-card {
                border: 2px solid gray !important;
                border-radius: 8px;
            }
            input[type = 'text'] { font-weight: bold; }
        """),

    ui.div("✨ TernTables ✨", class_="section-title"),

    ui.h5("Step 1: Upload File", class_="step-header"),
    ui.layout_columns(
        ui.card(".csv & .xlsx files are accepted. Please refresh page upon re-uploading a new file", ui.input_file("data_file", "", accept=[".csv", ".xlsx"]), width="100%"),
        ui.card("Example Output File", ui.download_button("download_example", "NOT IMPLEMENTED")),
        col_widths=(8, 4),
        ),
        
    ui.h5("Step 2: Select Variables", class_="step-header"),
    ui.output_ui('select_columns'),
    
    ui.h5("Step 3: Table Options", class_="step-header"),

    ui.layout_columns(    
        # Table Name
        ui.card(ui.input_text("table_name", "Input Table Name", placeholder="Enter table name",width="100%")),

        # Grouping Variable
        ui.card(ui.output_ui("grouping_variable")),
        
        col_widths= (6,6,)
        ),

    ui.layout_columns(    
        # Formatting Options
        ui.card(ui.input_numeric("decimals_table", "Table - # Decimals", 2, min=0, max=5)),
        ui.card(ui.input_numeric("decimals_pvalue", "P-Val - # Decimals", 3, min=0, max=5)),
        ui.card(ui.input_radio_buttons("output_format", "Output Format", ["n (%)", "% (n)"])),
        ui.card(ui.input_radio_buttons("show_odds_ratio", "Show Odds Ratio", ["No (Default)", "Yes"])),
        ui.card(ui.input_radio_buttons("remove_blanks", "Remove Unknown Values", ["No (Default)", "Yes"])),
        # ui.card(ui.input_radio_buttons("remove_blanks", ui.tags.span("Remove Unknown Values ",ui.tooltip(ui.icon("info-circle"),"Customize how each variable appears in the final table.")), ["No (Default)", "Yes"]),
        
        col_widths= (2,2,2,2,4)
        ),

    ui.h5("Step 4: Customize Table & Rows", class_="step-header"),
    ui.p("Update variable names, type, subheading and position. Variables that have the same position value will be ordered alphabetically. Variable types include:"),
    ui.layout_columns(    
        ui.card(ui.tags.strong("Categorical (Y/N): "), "Fisher's Exact Test", ui.tags.em("ex: Smoking, Diabetes")),
        ui.card(ui.tags.strong("Categorical (Dichotomous): "), "Fisher's Exact Test", ui.tags.em("ex: Sex")),
        ui.card(ui.tags.strong("Categorical (Multinomial): "), "\u03C72 test", ui.tags.em("ex: Race")),
        ui.card(ui.tags.strong("Ratio Continuous: "), "T-Test", ui.tags.em("ex: Age, GFR")),
        ui.card(ui.tags.strong("Ordinal Discrete: "), "Wilcoxon", ui.tags.em("ex: GCS, Tumor Grade")),
        ),

    # Subheadings
    ui.input_text("subheading_1", "Subheading 1", placeholder="Enter subheading 1 name"), 
    # ui.div("subheading_1", "Subheading 1"),
    ui.output_ui("var_settings_1"),

    ui.input_text("subheading_2", "Subheading 2", placeholder="Enter subheading 2 name"), 
    # ui.div("subheading_2", "Subheading 2"),
    ui.output_ui("var_settings_2"),
    
    ui.input_text("subheading_3", "Subheading 3", placeholder="Enter subheading 3 name"), 
    # ui.div("subheading_3", "Subheading 3"),
    ui.output_ui("var_settings_3"),
    
    ui.input_text("subheading_4", "Subheading 4", placeholder="Enter subheading 4 name"), 
    # ui.div("subheading_4", "Subheading 4"),
    ui.output_ui("var_settings_4"),
    
    # Variable Selection UI (dynamically generated)
    ui.output_ui("var_settings"),
    
    # Calculate
    ui.input_action_button("calculate", "Calculate"),
    
    # Download Button
    ui.download_button("download_table", "Download Table"),

    # Save Configuration
    ui.input_action_button("save_config", "Save Configuration"),
    ui.input_action_button("load_config", "Load Configuration"),
    
)

################################################################################
######################### Shiny App Server #####################################
################################################################################
def server(input, output, session):
    sheet_names = reactive.Value([])
    excel_trigger = reactive.Value(False)

    data = reactive.Value({})  # Store uploaded data
    cleaned_data = reactive.Value({})  # Store cleaned data
    selected_columns = reactive.Value([])  # Store selected columns
    var_config = reactive.Value({})  # Store variable settings dynamically
    group_var = reactive.Value(None)  # Store grouping variable
    prev_group_var = reactive.Value(None)
    subheadings = { # Reactive values to track column assignments per subheading
        "subheading_1": reactive.Value([]),
        "subheading_2": reactive.Value([]),
        "subheading_3": reactive.Value([]),
        "subheading_4": reactive.Value([])
    }
    subheading_names = { # Reactive values to track column assignments per subheading
        "subheading_1": reactive.Value("subheading_1"),
        "subheading_2": reactive.Value("subheading_2"),
        "subheading_3": reactive.Value("subheading_3"),
        "subheading_4": reactive.Value("subheading_4")
    } 

    @output
    @render.ui
    def select_columns():
        return ui.input_selectize("column_selectize", "Select desired variables below:",  
                {  "": {"":""} },  
                multiple=True,  
                width="100%",
            )  
       
    # def show_modal_on_excel():
    #     if len(sheet_names.get()) > 1 and excel_trigger.get():
    #         ui.modal_show(
    #             ui.modal(
    #                 ui.input_select("selected_sheet", "Select a Sheet", choices=sheet_names.get()),
    #                 title="Choose a Sheet",
    #                 easy_close=False,
    #                 footer=ui.modal_button("Confirm"),
    #             )
    #         )
    #         excel_trigger.set(False)  # Reset the trigger after showing the modal

    @reactive.effect
    def _():
        if input.data_file():
            file_info = input.data_file()[0]
            ext = os.path.splitext(file_info["name"])[-1]
            
            if ext == ".csv":
                df = pd.read_csv(file_info["datapath"])  # Reads header row by default
            elif ext == ".xlsx":
                with open(file_info["datapath"], "rb") as f:
                    xls = pd.ExcelFile(f)
                    sheet_names.set(xls.sheet_names)
                    print("Sheet Names", sheet_names.get)
                
                # excel_trigger.set(True)
                # show_modal_on_excel()

                selected = input.selected_sheet() or sheet_names.get()[0]  # Default to the first sheet if none selected
                df = pd.read_excel(file_info["datapath"], sheet_name=selected)

            # Clean column names: strip and remove non-alphanumeric chars
            clean_columns = [re.sub(r'\W+', '', col.strip()) for col in df.columns]
            df.columns = clean_columns
            df = df.apply(lambda x: x.str.strip() if x.dtype == "object" else x)

            data.set(df)  # Store data in reactive value
            clean_df = df.replace(missing_values, np.nan)
            cleaned_data.set(clean_df)  # Store cleaned data in reactive value
            
            column_dict = {col: col for col in df.columns}
        
            default_type = "Omit"
            default_position = 15
            
            # Store variable settings in a dictionary
            if not var_config.get():
                var_config.set({col: {
                    "type": default_type, 
                    "name": col, 
                    "subheading": "subheading_1",
                    "position": default_position,
                    "ref_val": None,
                } for col in df.columns})

            ui.update_selectize(  
                "column_selectize",  
                choices={"":column_dict}
            )  

    def column_selectize():
        available_columns = input.column_selectize()
        selected_columns.set(available_columns)

        all_assigned = set().union(*[set(subheadings[s]()) for s in subheadings])
        for col in available_columns:
            if col not in all_assigned:
                subheadings["subheading_1"].set(subheadings["subheading_1"]() + [col])

        # Update dropdown choices but don't set selected value here
        ui.update_select("grouping_var", choices=available_columns)

        try:
            df = data.get()
            for col in df.columns:
                if col not in available_columns:
                    for subheading in subheadings: # Remove from all subheadings if the column is not selected
                        if col in subheadings[subheading]():
                            subheadings[subheading].set([c for c in subheadings[subheading]() if c != col])
        except:
            pass

    @reactive.effect
    def sync_column_selection_with_subheadings():
        for subheading in subheadings:
            generate_subheading_ui(subheading)

    @reactive.effect
    def watch_column_changes():
        column_selectize()

    @reactive.effect
    def manage_group_var():
        available = input.column_selectize()
        selected = input.grouping_var()
        if not group_var.get() and available:
            print("Initializing group_var:", available[0])
            group_var.set(available[0])
            ui.update_select("grouping_var", choices=available, selected=available[0])
        elif selected and selected != group_var.get():
            group_var.set(selected)

    # Set Grouping Variable for analysis
    @output
    @render.ui
    def grouping_variable():
        return ui.input_select("grouping_var", "Grouping Variable (Table Column)", choices=[])

    # Update columns under subheadings
    def generate_subheading_ui(subheading_key):
        if input.remove_blanks() == "Yes":
            df = cleaned_data.get()
        else:
            df = data.get()

        if df is None or not isinstance(df, pd.DataFrame) or df.empty:  
            return 
        
        columns = subheadings[subheading_key]()
        if not columns:
            return ui.p("No variables assigned yet.")

        return ui.layout_columns(
        *[
            ui.card(
                ui.h5(col),
                ui.p(", ".join(map(str, df[col].dropna().unique()[:5]))),
                ui.input_text(
                    f"name_{col}",
                    "Variable Name",
                    value=var_config.get()[col]["name"]
                ),
                ui.input_select(
                    f"var_type_{col}",
                    "Variable Type",
                    variable_types,
                    selected=var_config.get()[col]["type"],
                ),
                ui.input_select(
                    f"subheading_{col}",
                    "Subheading",
                    [subheading_val.get() for subheading_val in subheading_names.values()],
                    selected=var_config.get()[col]["subheading"],
                ),
                ui.input_select(
                    f"position_{col}",
                    "Position",
                    list(range(1,31)),
                    selected=var_config.get()[col]["position"],
                ),
                ui.input_select(
                    f"ref_val_{col}",
                    "Reference Value (For Odds Ratio of Dichotomous Variables)",
                    df[col].unique().tolist(),
                ),
                class_="variable-card",
                id=f"{subheading_key}_{col}",
            )
            for col in columns
        ],
        col_widths=(4),
        # width=1, 
        class_="droppable-area",
        )
    
    @output
    @render.ui
    def var_settings_1():
        return generate_subheading_ui("subheading_1")

    @output
    @render.ui
    def var_settings_2():
        return generate_subheading_ui("subheading_2")

    @output
    @render.ui
    def var_settings_3():
        return generate_subheading_ui("subheading_3")

    @output
    @render.ui
    def var_settings_4():
        return generate_subheading_ui("subheading_4")

   
    # JavaScript to enable drag-and-drop using SortableJS
    ui.tags.script(
        """
        document.addEventListener("DOMContentLoaded", function() {
            document.querySelectorAll('.draggable-list').forEach(list => {
                new Sortable(list, {
                    group: 'shared',
                    animation: 150,
                    onEnd: function(evt) {
                        let movedVar = evt.item.dataset.var;
                        let newGroup = evt.to.id.replace('list-', '');
                        
                        // Update the server-side reactive variable
                        Shiny.setInputValue("dragged_var", JSON.stringify({movedVar, newGroup}));
                    }
                });
            });
        });
        """
    )
    # Handle drag-and-drop updates in the server
    @reactive.effect
    def update_subheadings():
        drag_event = input.dragged_var()
        if drag_event:
            drag_data = json.loads(drag_event)
            moved_var = drag_data["movedVar"]
            new_group = drag_data["newGroup"]

            # Remove from old subheading
            for key in subheadings:
                if moved_var in subheadings[key]():
                    subheadings[key].set([v for v in subheadings[key]() if v != moved_var])

            # Add to new subheading
            subheadings[new_group].set(subheadings[new_group]() + [moved_var])
            
    
    # Update variable settings dynamically when inputs change
    @reactive.effect
    def update_var_config():
        df = data.get()
        if df is None or not isinstance(df, pd.DataFrame) or df.empty:  
            return
        
        if selected_columns.get() is None or len(selected_columns.get()) == 0:
            return
        
        updated_config = var_config.get()

        print("Currently Selected Columns",selected_columns.get())
        
        # for col in df.columns:
        for col in selected_columns.get():
            new_subheading = input[f"subheading_{col}"]()
            old_subheading = var_config.get()[col]["subheading"]
            
            new_subheading_mapped = next((k for k, v in subheading_names.items() if v() == new_subheading), None)
            old_subheading_mapped = next((k for k, v in subheading_names.items() if v() == old_subheading), "Subheading 1")

            print(new_subheading_mapped, old_subheading_mapped)
            
            print("❗️ Updating configuration from...", updated_config[col])
            updated_config[col]["type"] = input[f"var_type_{col}"]()
            updated_config[col]["name"] = input[f"name_{col}"]() 
            updated_config[col]["position"] = int(input[f"position_{col}"]())
            updated_config[col]["subheading"] = input[f"subheading_{col}"]() 
            updated_config[col]["ref_val"] = input[f"ref_val_{col}"]()
            print("-                           to...", updated_config[col])
            
            # If the subheading has changed, move the column to the new subheading
            if new_subheading_mapped != old_subheading_mapped:
                # Remove the variable from the current subheading
                subheadings[old_subheading_mapped].set(
                    [c for c in subheadings[old_subheading_mapped]() if c != col]
                )
                
                # Add the variable to the new subheading
                subheadings[new_subheading_mapped].set(
                    subheadings[new_subheading_mapped]() + [col]
                )
                
                # Debugging print statement to track the change
                print(f"Moved {col} from {old_subheading_mapped} to {new_subheading_mapped}")

                generate_subheading_ui(new_subheading_mapped)
                generate_subheading_ui(old_subheading_mapped)

        print()
        var_config.set(updated_config)  # Update stored config

    @reactive.effect
    def update_subheading_names():
        for key in subheadings.keys(): # subheadings = {"subheading_1": ..., etc.}
            # input.get returns a Value instance; call .is_valid() and .__call__() to get safely
            try:
                input_val = input.get(key)
                if input_val is not None and input_val.is_valid():
                    text_input = input_val()
                    print("Subheading name:", key, "Text input:", text_input)
                    if text_input and text_input.strip() != "":
                        subheading_names[key].set(text_input.strip())
                        print("✅ Subheading name updated:", key, "→", text_input.strip())
            except Exception as e:
                print(f"⚠️ Failed to update subheading {key}: {e}")
                continue        


    # Perform statistical analysis when the "Calculate" button is clicked
    @reactive.effect
    @reactive.event(input.calculate)
    def calculate_statistical_analysis():
        print("🔄 Calculate button pressed. Updating variable configurations...")
        if input.remove_blanks() == "Yes":
            df = cleaned_data.get()
        else:
            df = data.get()

        if df is None or not isinstance(df, pd.DataFrame) or df.empty:  
            return
        
        try:
            curr_group_var = group_var.get()  # Get the selected grouping column
            decimals_pval = input.decimals_pvalue()
            decimals_tab = input.decimals_table()
            output_format = input.output_format()
    
            updated_config = var_config.get()
           
            # Perform statistical analysis using the grouping variable
            if len(selected_columns.get()) > 0:
                for col in df.columns:
                    if col != curr_group_var and col in selected_columns.get():
                        print(f"\n📂 Processing Variable: {col}", updated_config[col])
                        
                        var_type = updated_config[col]["type"]
                        
                        if var_type != "Omit":
                            p_value = run_statistical_test(df, curr_group_var, var_type, col, decimals_pval)
                            
                            # Store the p-value in the var_config dictionary
                            updated_config[col]["p_value"] = p_value
                            print(f"Column: {col}, Grouping Variable: {curr_group_var}, p-value: {p_value}")

                            # Perform aggregate analysis and update var_config with the results
                            aggregate_result = perform_aggregate_analysis(df, curr_group_var, var_type, col, decimals_tab, output_format, updated_config[col])
                            if aggregate_result:
                                updated_config[col].update(aggregate_result)
                            
                            print("After: ", updated_config[col])

                var_config.set(updated_config)

                print(var_config)

                ui.notification_show("✅ Calculation complete! File ready to download", duration=5, type="message")
        except:
            return

    # Download Button - Trigger to save table in .docx format
    # Updated download_table function
    @session.download()
    def download_table():
        # Retrieve the data and var_config
        if input.remove_blanks() == "Yes":
            df = cleaned_data.get()
        else:
            df = data.get()
        updated_config = var_config.get()
        
        if df is None or not isinstance(df, pd.DataFrame) or df.empty:  
            return None  # Return None if no data is available
        
        # Generate the Word table document
        doc_filename = create_word_table(df, updated_config, group_var.get(), subheadings, subheading_names, input.table_name(), input.show_odds_ratio(), input.output_format())  
        
        return doc_filename  # Return the Word document file for download

app = App(app_ui, server)